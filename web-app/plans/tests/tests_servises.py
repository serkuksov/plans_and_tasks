from datetime import datetime

from django.test import TestCase
from django.contrib.auth import get_user_model
from docx import Document
from docxcompose.composer import Composer

from ..servises import servises
from ..servises import export_in_doc
from .. import models
from accounts.models import UserDeteil, Division


class ServisesTestCase(TestCase):
    """Тесты сервисов"""
    def test_add_date_delta(self):
        date = datetime(day=1, month=1, year=2023)
        first = servises.add_date_delta(date=date, days=5, months=-1, years=-6)
        self.assertEqual(first, datetime(day=6, month=12, year=2016))

        date = datetime(day=1, month=1, year=2023)
        first = servises.add_date_delta(date=date, days=22, months=1, years=1)
        self.assertEqual(first, datetime(day=23, month=2, year=2024))

        date = datetime(day=1, month=1, year=2023)
        first = servises.add_date_delta(date=date, days=0, months=0, years=0)
        self.assertEqual(first, datetime(day=1, month=1, year=2023))

    def test_get_ofset_number(self):
        first = servises.get_ofset_number('+1')
        self.assertEqual(first, 1)

        first = servises.get_ofset_number('-3')
        self.assertEqual(first, -3)

        first = servises.get_ofset_number('2')
        self.assertEqual(first, 0)

    def test_get_ofset_params_date(self):
        first = servises.get_ofset_params_date(days='+1', months='5', years='-1')
        self.assertEqual(first, {'days': 1, 'months': 0, 'years': -1})

    def test_get_fixed_number(self):
        first = servises.get_fixed_number('+1')
        self.assertEqual(first, 0)

        first = servises.get_fixed_number('-3')
        self.assertEqual(first, 0)

        first = servises.get_fixed_number('2')
        self.assertEqual(first, 2)

    def test_set_fixed_args_date(self):
        date = datetime(day=1, month=1, year=2023)
        first = servises.set_fixed_args_date(date=date, days=-1, months=-2, years=-3)
        self.assertEqual(first, date)

        date = datetime(day=1, month=1, year=2023)
        first = servises.set_fixed_args_date(date=date, days=5, months=6, years=2025)
        self.assertEqual(first, datetime(day=5, month=6, year=2025))

        date = datetime(day=1, month=1, year=2023)
        first = servises.set_fixed_args_date(date=date, days=0, months=0, years=0)
        self.assertEqual(first, date)

        date = datetime(day=1, month=1, year=2023)
        first = servises.set_fixed_args_date(date=date, days=50, months=6, years=2025)
        self.assertEqual(first, datetime(day=30, month=6, year=2025))

    def test_get_fixed_params_date(self):
        first = servises.get_fixed_params_date(days='+1', months='5', years='-1')
        self.assertEqual(first, {'days': 0, 'months': 5, 'years': 0})

    def test_get_data_with_working_day(self):
        first = servises.get_data_with_working_day(datetime(day=23, month=2, year=2023))
        self.assertEqual(first, datetime(day=22, month=2, year=2023))

        first = servises.get_data_with_working_day(datetime(day=15, month=3, year=2023))
        self.assertEqual(first, datetime(day=15, month=3, year=2023))

        first = servises.get_data_with_working_day(datetime(day=19, month=3, year=2023))
        self.assertEqual(first, datetime(day=17, month=3, year=2023))


class ExportInDocTestCase(TestCase):
    """Тесты по созданию документа WORD на основе информации из БД"""
    def setUp(self):
        self.user = get_user_model().objects.create(username='user')
        self.division = Division.objects.create(name='division')
        UserDeteil.objects.create(
            user=self.user,
            second_name='second_name',
            phone_number=3399,
            division=self.division,
        )
        self.plan = models.Plan.objects.create(
            name='plan_name',
            description='plan_description',
            completion_date=datetime(day=1, month=1, year=2023),
            user_creator=self.user.userdeteil,
            user_updater=self.user.userdeteil,
        )
        self.perfomer = models.Perfomer.objects.create(division=self.division)
        self.taskgroup = models.TaskGroup.objects.create(name='taskgroup')
        self.pattern_plan = models.PatternPlan.objects.create(name='name_pattern_plan', description='description_pattern_plan')
        self.pattern_task = models.PatternTask.objects.create(
            pattern_plan=self.pattern_plan,
            task_group=self.taskgroup,
            name='name_pattern_task',
            divisin_perfomer=self.division,
        )
        self.task_1 = models.Task.objects.create(
            pattern_task=self.pattern_task,
            plan=self.plan,
            name='task_1_name',
            completion_date=datetime(day=1, month=1, year=2023),
            perfomer=self.perfomer,
            user_creator=self.user.userdeteil,
            user_updater=self.user.userdeteil,
        )
        self.task_2 = models.Task.objects.create(
            pattern_task=self.pattern_task,
            plan=self.plan,
            name='task_2_name',
            completion_date=datetime(day=1, month=1, year=2023),
            perfomer=self.perfomer,
            user_creator=self.user.userdeteil,
            user_updater=self.user.userdeteil,
            is_active=False,
        )

    def test_create_word_doc_for_plan(self):
        doc = export_in_doc.create_word_doc_for_plan(plan_id=self.plan.id)
        #doc.save('test.docx')
        self.assertIsInstance(doc, Composer)

    def test_get_composer_doc(self):
        composer_doc_none = export_in_doc._get_composer_doc([])
        self.assertIsNone(composer_doc_none)
        
        doc1 = Document()
        doc1.add_heading('test_header_1')
        doc2 = Document()
        doc2.add_heading('test_header_2')
        composer_doc = export_in_doc._get_composer_doc([doc1, doc2])
        text_composer_doc = [p.text for p in composer_doc.doc.paragraphs]
        text_test_doc = [p.text for p in Document('./plans/tests/files/test_composer_doc.docx').paragraphs]
        self.assertEqual(text_composer_doc, text_test_doc)

    def test_get_doc_for_plan(self):
        doc = export_in_doc._get_doc_for_plan(self.plan.id)
        test_doc = Document('./plans/tests/files/test_doc_for_plan.docx')
        text_paragraphs_doc = [p.text for p in doc.paragraphs]
        text_paragraphs_test_doc = [p.text for p in doc.paragraphs]
        text_tables_doc = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_tables_doc.append(cell.text)
        text_tables_test_doc = []
        for table in test_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_tables_test_doc.append(cell.text)
        self.assertEqual(text_paragraphs_doc, text_paragraphs_test_doc)
        self.assertEqual(text_tables_doc, text_tables_test_doc)

    def test_get_completion_date(self):
        completion_date = export_in_doc._get_completion_date(self.task_1)
        self.assertEqual(completion_date, str(datetime(day=1, month=1, year=2023).date().strftime('%d.%m.%Y')))

        completion_date = export_in_doc._get_completion_date(self.task_2)
        self.assertEqual(completion_date, 'Выполнено')