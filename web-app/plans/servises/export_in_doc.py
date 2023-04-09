from docx import Document
from docx.shared import Inches
from docxcompose.composer import Composer

from plans.models import Plan, Task, TaskGroup


def create_word_doc_for_plan(plan_id: str) -> Document:
    """Создает кончный документ План-графика в формате word получая id плана в БД"""
    start_doc = Document('./templates/doc/начало.docx')
    body_doc = _get_doc_for_plan(plan_id)
    end_doc = Document('./templates/doc/конец.docx')
    return _get_composer_doc(docs=[start_doc, body_doc, end_doc])


def _get_composer_doc(docs: list[Document,]) -> Document:
    """Собирает и отдает один документ word из списка документов"""
    composer_doc = Composer(docs[0])
    for doc in docs[1:]:
        composer_doc.append(doc)
    return composer_doc


def _get_doc_for_plan(plan_id: str) -> Document:
    """Формирует тело плана на основе шаблона"""
    doc = Document('./templates/doc/тело_документа.docx')
    #формирование заголовка плана
    header = Plan.objects.filter(id=plan_id).first().description
    doc.add_heading(header)
    #формирование таблицы задач плана
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Таблица'
    style_header_table = 'Заголовок таблицы'
    style_centr_table = 'Текст таблицы (по центру)'
    style_table_number_1 = 'Нумирация1'
    style_table_number_2 = 'Нумирация2'
    #формирование заголовков таблицы
    headers_table = ('№', 'Мероприятие', 'Ответственный исполнитель', 'Срок исполнения')
    hdr_cells = table.rows[0].cells
    for index, cell in enumerate(hdr_cells):
        cell.text = headers_table[index]
        cell.paragraphs[0].style = style_header_table
    hdr_cells[1].width = Inches(20)
    #формирование тела таблицы с разбивкой на группы задач
    task_groups = TaskGroup.objects.filter(patterntask__task__plan_id=plan_id).distinct()
    for task_group in task_groups:
        content_cells = table.add_row().cells
        content_cells[0].paragraphs[0].style = style_table_number_1
        merge_cells = content_cells[1].merge(content_cells[3])
        merge_cells.text = str(task_group.name)
        merge_cells.paragraphs[0].style = style_header_table
        tasks = (Task.objects.
                 select_related('perfomer__division').
                 filter(pattern_task__task_group=task_group, plan=plan_id))
        for task in tasks:
            content_cells = table.add_row().cells
            content_cells[0].paragraphs[0].style = style_table_number_2
            content_cells[1].text = str(task.name)
            content_cells[2].text = str(task.perfomer.division)
            content_cells[2].paragraphs[0].style = style_centr_table
            content_cells[3].text = _get_completion_date(task=task)
            content_cells[3].paragraphs[0].style = style_centr_table
    return doc


def _get_completion_date(task: Task) -> str:
    """Форматирование даты исполнения в зависимости от статуса задачи"""
    if task.is_active:
        return str(task.completion_date.strftime('%d.%m.%Y'))
    else:
        return 'Выполнено'
